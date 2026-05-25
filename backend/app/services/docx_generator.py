"""
Gerador de DOCX — produz o documento Word final a partir do ArtigoGerado.
Suporta formato Vancouver (padrão) e estrutura para TCC (ABNT em breve).
"""
import os
import asyncio
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.config import settings
from app.models.schemas import ArtigoGerado, CKO


def _set_margins(doc: Document, top=3, bottom=2, left=3, right=2):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def _add_page_number(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(10)


def _configurar_estilos(doc: Document):
    styles = doc.styles

    # Normal
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = Pt(18)  # 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading 1 — seções principais
    h1 = styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2 — subseções
    h2 = styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(12)
    h2.font.bold = False
    h2.font.italic = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(3)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text.upper() if level == 1 else text, level=level)
    return p


def _safe_str(value) -> str:
    """Converte qualquer valor para string de forma segura."""
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, dict):
        # dict inesperado — usa 'formatada' ou 'texto' se disponível
        return str(value.get("formatada") or value.get("texto") or "")
    return str(value)


def _add_paragraph(doc: Document, text, first_line_indent: bool = True):
    p = doc.add_paragraph(_safe_str(text))
    p.style = doc.styles["Normal"]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    return p


def _superscript(run):
    """Apply superscript formatting to a run via XML."""
    rPr = run._r.get_or_add_rPr()
    vertAlign = OxmlElement("w:vertAlign")
    vertAlign.set(qn("w:val"), "superscript")
    rPr.append(vertAlign)


def _add_title_page(doc: Document, artigo: ArtigoGerado, cko: CKO):
    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(artigo.titulo)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(6)

    # ── Autores ────────────────────────────────────────────────────────────────
    autores = cko.autores if cko.autores else []
    if autores:
        # Collect unique affiliations preserving order
        afils: list[str] = []
        for a in autores:
            inst = (a.instituicao or "").strip()
            if inst and inst not in afils:
                afils.append(inst)

        # Author byline: "Nome¹, Nome²..."
        byline_p = doc.add_paragraph()
        byline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        byline_p.paragraph_format.space_after = Pt(4)
        for idx, autor in enumerate(autores):
            if idx > 0:
                sep = byline_p.add_run(", ")
                sep.font.name = "Times New Roman"
                sep.font.size = Pt(11)
            name_run = byline_p.add_run(autor.nome)
            name_run.font.name = "Times New Roman"
            name_run.font.size = Pt(11)
            # Add superscript affiliation number if there are affiliations
            if afils and (autor.instituicao or "").strip():
                sup_num = afils.index((autor.instituicao or "").strip()) + 1
                sup_run = byline_p.add_run(str(sup_num))
                sup_run.font.name = "Times New Roman"
                sup_run.font.size = Pt(9)
                _superscript(sup_run)
            # Mark corresponding author with asterisk
            if autor.correspondente:
                cor_run = byline_p.add_run("*")
                cor_run.font.name = "Times New Roman"
                cor_run.font.size = Pt(9)
                _superscript(cor_run)

        # Affiliations
        for i, inst in enumerate(afils, start=1):
            afil_p = doc.add_paragraph()
            afil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            afil_p.paragraph_format.space_after = Pt(2)
            num_run = afil_p.add_run(str(i))
            num_run.font.name = "Times New Roman"
            num_run.font.size = Pt(9)
            _superscript(num_run)
            inst_run = afil_p.add_run(" " + inst)
            inst_run.font.name = "Times New Roman"
            inst_run.font.size = Pt(10)

        # Corresponding author line
        corresp = next((a for a in autores if a.correspondente), None)
        if corresp and corresp.email:
            cor_p = doc.add_paragraph()
            cor_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cor_p.paragraph_format.space_after = Pt(2)
            label_run = cor_p.add_run("*Autor correspondente: ")
            label_run.font.name = "Times New Roman"
            label_run.font.size = Pt(10)
            label_run.font.italic = True
            email_run = cor_p.add_run(corresp.email)
            email_run.font.name = "Times New Roman"
            email_run.font.size = Pt(10)
            email_run.font.italic = True

        # ORCID line (for those who provided it)
        orcid_autores = [(a.nome, a.orcid) for a in autores if a.orcid]
        if orcid_autores:
            orc_p = doc.add_paragraph()
            orc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            orc_p.paragraph_format.space_before = Pt(2)
            orc_p.paragraph_format.space_after = Pt(4)
            orc_run = orc_p.add_run("ORCID: " + " | ".join(f"{n}: {o}" for n, o in orcid_autores))
            orc_run.font.name = "Times New Roman"
            orc_run.font.size = Pt(9)
            orc_run.font.color.rgb = RGBColor(80, 80, 80)

    # Palavras-chave (garante que cada item é string)
    doc.add_paragraph()
    kw_para = doc.add_paragraph()
    kw_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw_run = kw_para.add_run("Palavras-chave: ")
    kw_run.font.bold = True
    kw_run.font.size = Pt(12)
    kw_run.font.name = "Times New Roman"
    kws = "; ".join(_safe_str(k) for k in artigo.palavras_chave if k) + "."
    kw_val = kw_para.add_run(kws)
    kw_val.font.size = Pt(12)
    kw_val.font.name = "Times New Roman"

    doc.add_paragraph()


def _add_resumo(doc: Document, artigo: ArtigoGerado):
    _add_heading(doc, "Resumo")
    res = artigo.resumo

    for label, texto in [
        ("Introdução", res.introducao),
        ("Apresentação do Caso", res.caso),
        ("Discussão", res.discussao),
        ("Conclusão", res.conclusao),
    ]:
        texto_str = _safe_str(texto)
        if not texto_str:
            continue  # pula subseções vazias em vez de gerar linha em branco
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = Pt(12)  # simples
        bold_run = p.add_run(f"{label}: ")
        bold_run.font.bold = True
        bold_run.font.name = "Times New Roman"
        bold_run.font.size = Pt(12)
        text_run = p.add_run(texto_str)
        text_run.font.name = "Times New Roman"
        text_run.font.size = Pt(12)

    doc.add_paragraph()


def _add_secao(doc: Document, titulo: str, paragrafos: list[str]):
    _add_heading(doc, titulo)
    for texto in paragrafos:
        _add_paragraph(doc, texto)
    doc.add_paragraph()


def _add_referencias(doc: Document, artigo: ArtigoGerado):
    import re
    _add_heading(doc, "Referências")
    for ref in artigo.referencias:
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = Pt(12)  # simples
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)

        # Avoid duplicate numbering: if `formatada` already starts with
        # a number followed by a dot (e.g. "1. Marx RE..."), do NOT
        # prepend the number again.
        formatada = ref.formatada or ""
        if re.match(r"^\d+\.", formatada.lstrip()):
            texto_ref = formatada
        else:
            texto_ref = f"{ref.numero}. {formatada}"

        run = p.add_run(texto_ref)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        # DOI handling
        doi = ref.doi if ref.doi else ""
        doi = doi.strip()
        if doi:
            # Normalise: ensure it is displayed as a full URL
            if not doi.startswith("https://"):
                doi_url = f"https://doi.org/{doi}"
            else:
                doi_url = doi
            doi_run = p.add_run(f" doi: {doi_url}")
            doi_run.font.name = "Times New Roman"
            doi_run.font.size = Pt(12)
        else:
            # No DOI — add a red warning
            no_doi_run = p.add_run(" [DOI não localizado — verifique manualmente]")
            no_doi_run.font.name = "Times New Roman"
            no_doi_run.font.size = Pt(12)
            no_doi_run.font.color.rgb = RGBColor(204, 0, 0)


def _add_disclaimer(doc: Document):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(
        "AVISO: Este documento foi gerado com auxílio de inteligência artificial (RCCS v1.0). "
        "O autor é integralmente responsável pela revisão do conteúdo antes de qualquer submissão. "
        "A plataforma não garante aceitação em periódicos científicos."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)


def _add_imagens(doc: Document, cko: CKO, sessao_id: str):
    """Insert figures after Caso Clínico with CARE-compliant captions."""
    if not cko.imagens:
        return

    imagens_sorted = sorted(cko.imagens, key=lambda x: x.numero_figura)
    images_dir = Path("images_output") / sessao_id

    for img in imagens_sorted:
        if not img.filename:
            continue
        img_path = images_dir / img.filename
        if not img_path.exists():
            continue

        try:
            # Insert image centered
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img.add_run()
            run.add_picture(str(img_path), width=Inches(5.0))

            # CARE caption: "Figura X – legenda."
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(12)
            run_cap = p_cap.add_run(f"Figura {img.numero_figura} \u2013 {img.legenda}")
            run_cap.font.name = "Times New Roman"
            run_cap.font.size = Pt(11)
            run_cap.font.italic = True
        except Exception as exc:
            # A missing or corrupt image must not break the whole document
            import logging
            logging.getLogger(__name__).warning(
                "Falha ao inserir figura %s: %s", img.filename, exc
            )


async def gerar_docx_bytes(artigo: ArtigoGerado, cko: CKO, sessao_id: str) -> bytes:
    """Gera o DOCX em memória e retorna os bytes. Não depende de filesystem."""
    import io
    loop = asyncio.get_event_loop()
    buf = io.BytesIO()
    await loop.run_in_executor(None, _build_docx_to_buffer, buf, artigo, cko, sessao_id)
    return buf.getvalue()


# Mantido por compatibilidade — internamente usa gerar_docx_bytes
async def gerar_docx(sessao_id: str, artigo: ArtigoGerado, cko: CKO) -> str:
    import tempfile, io
    data = await gerar_docx_bytes(artigo, cko, sessao_id)
    output_dir = Path(settings.docx_output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    filepath = output_dir / f"{sessao_id}.docx"
    filepath.write_bytes(data)
    return str(filepath)


def _build_docx_to_buffer(buf, artigo: ArtigoGerado, cko: CKO, sessao_id: str):
    """Constrói o Document e salva no buffer (BytesIO ou Path)."""
    doc = Document()
    _set_margins(doc)
    _configurar_estilos(doc)
    _add_page_number(doc)

    _add_title_page(doc, artigo, cko)
    _add_resumo(doc, artigo)
    _add_secao(doc, "Introdução", artigo.introducao)
    _add_secao(doc, "Apresentação do Caso", artigo.caso_clinico)
    _add_imagens(doc, cko, sessao_id)  # figures after case presentation
    _add_secao(doc, "Discussão", artigo.discussao)
    _add_secao(doc, "Conclusão", artigo.conclusao)
    _add_referencias(doc, artigo)
    _add_disclaimer(doc)

    doc.save(buf)


# Alias para compatibilidade com código legado
def _build_docx(filepath: Path, artigo: ArtigoGerado, cko: CKO, sessao_id: str):
    _build_docx_to_buffer(filepath, artigo, cko, sessao_id)
